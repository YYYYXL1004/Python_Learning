import os
import re

from docx import Document


# Word文件保存路径
SAVE_DIR = os.path.join(os.path.dirname(__file__), "exp3_files")
DOCX_FILE = os.path.join(SAVE_DIR, "article.docx")


def make_test_docx():
    """生成一个用于测试的Word文章。"""
    if not os.path.exists(SAVE_DIR):
        os.mkdir(SAVE_DIR)

    doc = Document()
    doc.add_heading("AABB词语测试文章", level=1)
    doc.add_paragraph("学习Python要踏踏实实，不能马马虎虎。")
    doc.add_paragraph("写程序时要简简单单地表达思路，也要时时刻刻注意细节。")
    doc.add_paragraph("春天来了，花花草草都变得密密麻麻，校园里热热闹闹。")
    doc.add_paragraph("这个段落中没有符合要求的词语，用来测试程序。")
    doc.save(DOCX_FILE)


def read_docx_text(file_path):
    """读取Word文档中所有段落的文本。"""
    doc = Document(file_path)

    # 使用列表推导式取出所有非空段落
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def find_aabb_words(text):
    """查找AABB形式的中文词语。"""
    # 常用汉字Unicode编码范围大致为：\u4e00-\u9fff
    # ([一-龥])\1 表示第一个汉字连续出现两次
    # ([一-龥])\2 表示第二个汉字连续出现两次
    pattern = r"([\u4e00-\u9fff])\1(?!\1)([\u4e00-\u9fff])\2"
    result = []

    for word in (m.group() for m in re.finditer(pattern, text)):
        # 保持原文出现顺序，同时去掉重复词语
        if word not in result:
            result.append(word)

    return result


def main():
    if not os.path.exists(DOCX_FILE):
        make_test_docx()

    text = read_docx_text(DOCX_FILE)
    words = find_aabb_words(text)

    print("文章中的AABB形式词语如下：")
    for word in words:
        print(word)


if __name__ == "__main__":
    main()
