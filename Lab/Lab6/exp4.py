import os

from docx import Document
from docx.shared import RGBColor


# Word文件保存路径
SAVE_DIR = os.path.join(os.path.dirname(__file__), "exp4_files")
DOCX_FILE = os.path.join(SAVE_DIR, "article.docx")


def make_test_docx():
    """生成一个带有不同格式文本的Word文档。"""
    if not os.path.exists(SAVE_DIR):
        os.mkdir(SAVE_DIR)

    doc = Document()
    doc.add_heading("Word Run格式测试", level=1)

    p = doc.add_paragraph()
    p.add_run("这是一段普通文本，")

    run = p.add_run("红色文本")
    run.font.color.rgb = RGBColor(255, 0, 0)

    p.add_run("，这里是")

    run = p.add_run("加粗文本")
    run.bold = True

    p.add_run("，最后是")

    run = p.add_run("红色加粗文本")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    p.add_run("。")

    p = doc.add_paragraph()
    p.add_run("第二段中，")

    run = p.add_run("重点内容")
    run.bold = True

    p.add_run("需要认真阅读，")

    run = p.add_run("警告内容")
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.bold = True
    p.add_run("需要特别注意。")

    doc.save(DOCX_FILE)


def is_red(run):
    """判断run中的文本是否为红色。"""
    color = run.font.color
    if color is None or color.rgb is None:
        return False

    return color.rgb == RGBColor(255, 0, 0)


def get_special_text(file_path):
    """读取Word文件，找出红色文本、加粗文本以及红色且加粗的文本。"""
    doc = Document(file_path)

    result = {
        "红色文本": [],
        "加粗文本": [],
        "红色且加粗文本": []
    }

    # Word段落由多个run组成，每个run可以有自己的字体格式
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text.strip()
            if not text:
                continue

            red = is_red(run)
            bold = run.bold is True

            if red:
                result["红色文本"].append(text)
            if bold:
                result["加粗文本"].append(text)
            if red and bold:
                result["红色且加粗文本"].append(text)

    return result


def remove_duplicate(words):
    """使用集合去重，同时保留文本原来的出现顺序。"""
    used = set()
    new_words = []

    for word in words:
        if word not in used:
            new_words.append(word)
            used.add(word)

    return new_words


def print_result(result):
    """输出查找结果。"""
    for title, words in result.items():
        print(title + "：")

        # 使用列表保存去重后的文本
        words = remove_duplicate(words)
        if words:
            for word in words:
                print(word)
        else:
            print("无")

        print()


def main():
    if not os.path.exists(DOCX_FILE):
        make_test_docx()

    result = get_special_text(DOCX_FILE)
    print_result(result)


if __name__ == "__main__":
    main()
